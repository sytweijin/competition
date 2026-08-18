"""Large-project mode: planner branch, deterministic fallback, and volunteer slots."""

from datetime import date

import pytest

from app.agents.scoring import assign_with_balance
from app.coordinator import Coordinator
from app.models.schemas import (
    AssignmentInput, CourseInfo, DraftOperation, FullPlan,
    ManualAssignmentRequest, PlanOutput, ProjectModule, QAOutput, ReportOutput,
    SubTask, TeamMember, TimelineOutput, Volunteer,
)
from app.services.project_service import (
    ProjectServiceError, apply_manual_assignment, generate_draft,
    mutate_draft, update_volunteer_pool,
)


def _large_input(project_mode: str = "large_project") -> AssignmentInput:
    return AssignmentInput(
        course=CourseInfo(name="Campus Event", description="Plan and run a campus event"),
        members=[
            TeamMember(name="Alice", skill_tags=["planning"]),
            TeamMember(name="Bob", skill_tags=["execution"]),
        ],
        deadline=date(2026, 8, 20),
        project_mode=project_mode,
    )


def test_small_group_does_not_keep_member_hierarchy():
    data = AssignmentInput(
        course=CourseInfo(name="Small", description="No hierarchy"),
        members=[TeamMember(name="Alice", manager="Bob")],
        deadline=date(2026, 8, 20),
        project_mode="small_group",
    )
    assert data.members[0].manager == ""


def test_large_project_fallback_marks_volunteer_demand():
    plan = Coordinator._fallback_large_project_plan(_large_input(), "LLM down")
    assert len(plan.modules) == 4
    assert len(plan.tasks) == 8
    assert all(task.module_id for task in plan.tasks)
    assert all(task.extra_helpers_needed >= 0 for task in plan.tasks)
    assert any(task.extra_helpers_needed > 0 for task in plan.tasks)
    assert all(
        task.suggested_people == 1 + task.extra_helpers_needed
        for task in plan.tasks
    )
    assert plan.tasks[0].assignee_id in {"Alice", "Bob"}


def test_large_project_quick_draft_uses_large_fallback():
    plan = generate_draft(_large_input(), use_ai=False)
    assert len(plan.modules) == 4
    assert len(plan.tasks) == 8
    assert all(task.module_id for task in plan.tasks)
    assert any(task.extra_helpers_needed > 0 for task in plan.tasks)


def test_large_project_draft_allows_empty_members():
    inp = _large_input()
    inp.members = []
    plan = generate_draft(inp, use_ai=False)
    assert len(plan.modules) == 4
    assert len(plan.tasks) == 8
    assert all(task.module_id for task in plan.tasks)
    assert all(task.assignee_id is None for task in plan.tasks)


def test_mutate_draft_supports_module_lifecycle():
    plan = generate_draft(_large_input(), use_ai=False)

    plan = mutate_draft(plan, [DraftOperation(op="add_module")])
    assert [m.id for m in plan.modules] == ["M1", "M2", "M3", "M4", "M5"]
    assert plan.modules[-1].order == 5

    plan = mutate_draft(plan, [DraftOperation(
        op="reorder_modules",
        ordered_module_ids=["M5", "M1", "M2", "M3", "M4"])])
    assert [m.id for m in plan.modules] == ["M5", "M1", "M2", "M3", "M4"]
    assert [m.order for m in plan.modules] == [1, 2, 3, 4, 5]

    plan = mutate_draft(plan, [DraftOperation(
        op="update_module", module_id="M5",
        module=ProjectModule(
            id="M5", name="宣传推广", description="面向目标用户",
            assignee_id="Alice"))])
    m5 = next(m for m in plan.modules if m.id == "M5")
    assert m5.name == "宣传推广"
    assert m5.assignee_id == "Alice"
    assert m5.order == 1

    plan = mutate_draft(plan, [DraftOperation(
        op="remove_module", module_id="M5")])
    assert [m.id for m in plan.modules] == ["M1", "M2", "M3", "M4"]

    plan = mutate_draft(plan, [DraftOperation(
        op="add", module_id="M2")])
    added = max(plan.tasks, key=lambda task: task.order)
    assert added.module_id == "M2"


def test_apply_manual_assignment_persists_module_assignees():
    fp = _large_full_plan()
    plan = fp.plan.model_copy(update={
        "tasks": [
            task.model_copy(update={"assignee_id": None})
            for task in fp.plan.tasks
        ],
        "modules": [
            module.model_copy(update={"assignee_id": None})
            for module in fp.plan.modules
        ],
    })
    fp = fp.model_copy(update={"plan": plan})

    updated = apply_manual_assignment(ManualAssignmentRequest(
        plan=fp,
        module_assignees={"M1": "Bob"},
    ))
    m1 = next(module for module in updated.plan.modules if module.id == "M1")
    assert m1.assignee_id == "Bob"
    assert all(
        task.assignee_id == "Bob"
        for task in updated.plan.tasks
        if task.module_id == "M1" and task.status != "completed"
    )


def test_legacy_volunteer_role_member_is_not_used_as_module_owner():
    fp = _large_full_plan()
    volunteer = TeamMember(
        name="Volunteer A", role="志愿者 / 外部协作者",
        skill_tags=[], daily_available_hours=4,
    )
    fp = fp.model_copy(update={
        "input": fp.input.model_copy(update={
            "members": [*fp.input.members, volunteer],
        }),
    })
    with pytest.raises(ProjectServiceError):
        apply_manual_assignment(ManualAssignmentRequest(
            plan=fp,
            module_assignees={"M1": "Volunteer A"},
        ))


def test_apply_manual_assignment_rejects_unknown_module():
    fp = _large_full_plan()
    with pytest.raises(ProjectServiceError, match="模块不存在"):
        apply_manual_assignment(ManualAssignmentRequest(
            plan=fp,
            module_assignees={"M99": "Alice"},
        ))


def test_large_project_draft_uses_planner_and_keeps_assessment(monkeypatch):
    from app.llm.client import LLMClient

    class FakeLLM:
        def chat_structured(self, **kwargs):
            return PlanOutput(
                tasks=[
                    SubTask(
                        id="T1", name="Planning", description="",
                        estimated_hours=4, required_skills=["planning"],
                        assignee_id="Alice", extra_helpers_needed=2,
                        suggested_people=3,
                    ),
                    SubTask(
                        id="T2", name="Execution", description="",
                        estimated_hours=8, dependencies=["T1"],
                        required_skills=["execution"], assignee_id="Bob",
                        extra_helpers_needed=0, suggested_people=1,
                    ),
                ],
                summary="Large project plan",
                member_assessment={"Alice": "Good at planning",
                                   "Bob": "Good at execution"},
            )

    monkeypatch.setattr(LLMClient, "get_shared", lambda: FakeLLM())
    plan = Coordinator().draft(_large_input())
    assert plan.tasks[0].extra_helpers_needed == 2
    assert plan.tasks[0].suggested_people == 3
    assert plan.member_assessment["Alice"] == "Good at planning"


def test_volunteer_demand_does_not_create_internal_collaborators():
    plan = PlanOutput(
        tasks=[
            SubTask(
                id="T1", name="Large task", description="",
                estimated_hours=6, required_skills=["execution"],
                assignee_id="Alice", extra_helpers_needed=2,
                suggested_people=3,
            )
        ],
        summary="one task",
    )
    members = [
        TeamMember(name="Alice", skill_tags=["execution"]),
        TeamMember(name="Bob", skill_tags=["execution"]),
    ]
    qa = assign_with_balance(plan, members)
    assignment = qa.assignments[0]
    assert assignment.presenter == "Alice"
    assert not assignment.qa_primary
    assert not assignment.qa_support


def test_small_group_suggestion_still_creates_collaborators():
    plan = PlanOutput(
        tasks=[
            SubTask(
                id="T1", name="Group task", description="",
                estimated_hours=4, required_skills=["writing"],
                assignee_id="Alice", suggested_people=3,
            )
        ],
        summary="one task",
    )
    members = [
        TeamMember(name="Alice", skill_tags=["writing"]),
        TeamMember(name="Bob", skill_tags=["writing"]),
        TeamMember(name="Carol", skill_tags=["writing"]),
    ]
    qa = assign_with_balance(plan, members)
    assignment = qa.assignments[0]
    assert assignment.qa_primary
    assert len(assignment.qa_support) >= 1


def _large_full_plan(project_mode: str = "large_project") -> FullPlan:
    """构造一个带志愿者需求任务的最小 FullPlan，供招募池测试使用。"""
    inp = _large_input(project_mode)
    plan = Coordinator._fallback_large_project_plan(inp, "test")
    return FullPlan(
        input=inp,
        plan=plan,
        timeline=TimelineOutput(tasks=[], critical_path=[], total_days=1, note=""),
        qa_matrix=QAOutput(assignments=[], workload={}, note=""),
        report=ReportOutput(summary="", risk_note=""),
    )


def _need_task_id(fp: FullPlan) -> str:
    return next(t.id for t in fp.plan.tasks if t.extra_helpers_needed > 0)


def _removable_need_task_id(fp: FullPlan) -> str:
    depended_on = {d for t in fp.plan.tasks for d in t.dependencies}
    return next(
        t.id for t in fp.plan.tasks
        if t.extra_helpers_needed > 0 and t.id not in depended_on
    )


def test_update_volunteer_pool_upsert_and_preserves_immutability():
    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    volunteers = [
        Volunteer(name="V1", task_id=task_id, status="待确认", contact="wx"),
        Volunteer(name="V2", task_id=task_id, status="已确认", contact="phone"),
    ]
    updated = update_volunteer_pool(fp, volunteers)
    assert [v.name for v in updated.volunteer_pool] == ["V1", "V2"]
    assert updated.volunteer_pool[0].contact == "wx"
    # 原计划不应被原地修改
    assert fp.volunteer_pool == []


def test_volunteer_pool_rejects_small_group_mode():
    fp = _large_full_plan(project_mode="small_group")
    with pytest.raises(ProjectServiceError, match="仅适用于大型项目模式"):
        update_volunteer_pool(fp, [])


def test_volunteer_pool_rejects_member_name_collision():
    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    with pytest.raises(ProjectServiceError, match="不能与团队成员重复"):
        update_volunteer_pool(fp, [Volunteer(name="Alice", task_id=task_id)])


def test_volunteer_pool_rejects_duplicate_names():
    fp = _large_full_plan()
    task_ids = [t.id for t in fp.plan.tasks if t.extra_helpers_needed > 0]
    if len(task_ids) < 2:
        task_ids = task_ids + task_ids
    with pytest.raises(ProjectServiceError, match="姓名重复"):
        update_volunteer_pool(fp, [
            Volunteer(name="V1", task_id=task_ids[0]),
            Volunteer(name="V1", task_id=task_ids[1]),
        ])


def test_volunteer_pool_rejects_missing_task():
    fp = _large_full_plan()
    with pytest.raises(ProjectServiceError, match="任务不存在"):
        update_volunteer_pool(fp, [Volunteer(name="V1", task_id="T99")])


def test_volunteer_pool_rejects_task_without_demand():
    fp = _large_full_plan()
    task_id = next(t.id for t in fp.plan.tasks if t.extra_helpers_needed == 0)
    with pytest.raises(ProjectServiceError, match="不需要招募志愿者"):
        update_volunteer_pool(fp, [Volunteer(name="V1", task_id=task_id)])


def test_volunteer_pool_rejects_quota_overflow():
    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    need = next(t for t in fp.plan.tasks if t.id == task_id).extra_helpers_needed
    volunteers = [
        Volunteer(name=f"V{i}", task_id=task_id, status="已确认")
        for i in range(need + 1)
    ]
    with pytest.raises(ProjectServiceError, match="超过需求"):
        update_volunteer_pool(fp, volunteers)


def test_volunteer_pool_allows_declined_beyond_quota():
    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    need = next(t for t in fp.plan.tasks if t.id == task_id).extra_helpers_needed
    volunteers = [
        Volunteer(name=f"V{i}", task_id=task_id, status="已确认")
        for i in range(need)
    ]
    volunteers.append(Volunteer(name="Backup", task_id=task_id, status="已婉拒"))
    updated = update_volunteer_pool(fp, volunteers)
    assert len(updated.volunteer_pool) == need + 1


def test_volunteer_pool_rejects_bad_status():
    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    with pytest.raises(ProjectServiceError, match="状态不合法"):
        update_volunteer_pool(fp, [Volunteer(name="V1", task_id=task_id, status="审核中")])


def test_plan_markdown_export_includes_volunteer_section():
    from app.web.routes import _plan_to_markdown

    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    fp = update_volunteer_pool(fp, [
        Volunteer(name="V1", task_id=task_id, status="已确认", contact="wx:123"),
    ])
    md = _plan_to_markdown(fp.model_dump())
    assert "志愿者招募计划" in md
    assert "认领明细" in md
    assert "V1" in md
    assert "wx:123" in md


def test_manual_assignment_preserves_volunteer_pool():
    from app.models.schemas import ManualAssignmentRequest
    from app.services.project_service import apply_manual_assignment

    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    fp = update_volunteer_pool(fp, [
        Volunteer(name="V1", task_id=task_id, status="已确认"),
    ])
    updated = apply_manual_assignment(ManualAssignmentRequest(plan=fp))
    assert [v.name for v in updated.volunteer_pool] == ["V1"]
    assert updated.version == fp.version


def test_edit_plan_cleans_volunteers_of_removed_task():
    from app.editor import edit_plan
    from app.models.schemas import EditPlanRequest, TaskEdit

    fp = _large_full_plan()
    task_id = _removable_need_task_id(fp)
    fp = update_volunteer_pool(fp, [
        Volunteer(name="V1", task_id=task_id, status="已确认"),
    ])
    updated = edit_plan(EditPlanRequest(
        plan=fp,
        edits=[TaskEdit(op="remove", task_id=task_id)],
    ))
    assert updated.volunteer_pool == []
    assert task_id not in {t.id for t in updated.plan.tasks}


def test_docx_export_contains_volunteer_section():
    import io

    from docx import Document

    from app.web.exporters import plan_to_docx

    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    fp = update_volunteer_pool(fp, [
        Volunteer(name="V1", task_id=task_id, status="已确认", contact="wx:123"),
    ])
    doc = Document(io.BytesIO(plan_to_docx(fp)))
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in doc.tables
        for row in table.rows for cell in row.cells)
    assert "志愿者招募计划" in text
    assert "V1" in text
    assert "wx:123" in text


def test_pdf_export_builds_with_volunteers():
    from app.web.exporters import plan_to_pdf

    fp = _large_full_plan()
    task_id = _need_task_id(fp)
    fp = update_volunteer_pool(fp, [
        Volunteer(name="V1", task_id=task_id, status="已确认"),
    ])
    data = plan_to_pdf(fp)
    assert len(data) > 1000
