"""
计划导出工具：将 FullPlan 转换为 Word(.docx) 和 PDF 格式。
供 /api/export 端点调用，提供普通用户友好的可打印文档。
"""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import xml.sax.saxutils as saxutils

from app.models.schemas import FullPlan


def _md_split_inline(text):
    """把一行里的 **bold** / *italic* 拆成 (text, bold, italic) 片段。"""
    import re
    parts = []
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*", text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False))
        if m.group(1) is not None:
            parts.append((m.group(1), True, False))
        else:
            parts.append((m.group(2), False, True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False))
    return parts or [(text, False, False)]


def _md_blocks(text):
    """把 markdown 文本切成 (kind, content) 块：h2/h3/li/p/table。"""
    if not text:
        return []
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # 表格
        if line.strip().startswith("|") and line.strip().endswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # 分隔行（全是 -/:/空格）跳过，不当作数据行
                if cells and all(set(c) <= set("-: ") for c in cells if c):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue
        # 标题
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h2", line[2:].strip()))
        elif line.startswith("- ") or line.startswith("* "):
            blocks.append(("li", line[2:].strip()))
        else:
            blocks.append(("p", line.strip()))
        i += 1
    return blocks


def _md_to_docx(doc, text, base_heading_level=1):
    """把 markdown 文本渲染进 Word 文档（标题/列表/粗体斜体/表格）。"""
    for kind, content in _md_blocks(text):
        if kind == "h2":
            doc.add_heading(content, level=base_heading_level)
        elif kind == "h3":
            doc.add_heading(content, level=base_heading_level + 1)
        elif kind == "li":
            p = doc.add_paragraph(style="List Bullet")
            for txt, bold, italic in _md_split_inline(content):
                if not txt:
                    continue
                r = p.add_run(txt)
                r.bold = bold or None
                r.italic = italic or None
        elif kind == "table":
            tbl = doc.add_table(rows=0, cols=len(content[0]), style="Table Grid")
            for ridx, row in enumerate(content):
                cells = tbl.add_row().cells
                for cidx, cell in enumerate(row):
                    if cidx < len(cells):
                        cells[cidx].text = cell
                        if ridx == 0:
                            for run in cells[cidx].paragraphs[0].runs:
                                run.bold = True
        else:  # p
            p = doc.add_paragraph()
            for txt, bold, italic in _md_split_inline(content):
                if not txt:
                    continue
                r = p.add_run(txt)
                r.bold = bold or None
                r.italic = italic or None


def _md_inline_to_pdf_html(text):
    """行内 markdown -> reportlab 可识别的 mini-HTML（先转义，再加 <b>/<i>）。"""
    import re
    out = []
    for txt, bold, italic in _md_split_inline(text):
        seg = saxutils.escape(txt)
        if bold:
            seg = "<b>" + seg + "</b>"
        if italic:
            seg = "<i>" + seg + "</i>"
        out.append(seg)
    return "".join(out)


def _md_to_pdf_story(text, body_style, h2_style, h3_style, font_name):
    """把 markdown 文本渲染成 reportlab flowable 列表。"""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    story = []
    for kind, content in _md_blocks(text):
        if kind == "h2":
            story.append(Paragraph(saxutils.escape(content), h2_style))
        elif kind == "h3":
            story.append(Paragraph(saxutils.escape(content), h3_style))
        elif kind == "li":
            story.append(Paragraph("• " + _md_inline_to_pdf_html(content), body_style))
        elif kind == "table":
            n = max(len(r) for r in content)
            cw = 17 * cm / n
            wrapped = [[Paragraph(str(cell), body_style) for cell in row] for row in content]
            tbl = Table(wrapped, colWidths=[cw] * n, repeatRows=1)
            st = TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4f46e5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#f1f5f9")]),
            ])
            if font_name:
                st.add("FONTNAME", (0, 0), (-1, -1), font_name)
            tbl.setStyle(st)
            story.append(tbl)
        else:  # p
            story.append(Paragraph(_md_inline_to_pdf_html(content), body_style))
        story.append(Spacer(1, 4))
    return story


def _register_cjk_font():
    """注册中文字体（PDF 默认不支持中文），失败则回退 Helvetica。

    支持跨平台：Windows（微软雅黑/宋体/黑体）、Linux（Noto/WenQuanYi）、
    macOS（PingFang/STHeiti），并允许通过 CJK_FONT_PATH 环境变量指定。
    """
    import os
    # 环境变量优先：允许显式指定字体文件
    env_path = os.getenv("CJK_FONT_PATH")
    if env_path:
        try:
            pdfmetrics.registerFont(TTFont("CJK", env_path))
            return "CJK"
        except Exception:
            pass
    # 跨平台候选字体
    candidates = [
        # Windows
        ("C:\\Windows\\Fonts\\msyh.ttc", "MSYH"),
        ("C:\\Windows\\Fonts\\simsun.ttc", "SimSun"),
        ("C:\\Windows\\Fonts\\simhei.ttf", "SimHei"),
        # Linux（常见中文字体包）
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoCJK"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoCJK"),
        ("/usr/share/fonts/wenquanyi/wqy-microhei/wqy-microhei.ttc", "WQY"),
        ("/usr/share/fonts/wqy-microhei/wqy-microhei.ttc", "WQY"),
        # macOS
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ("/System/Library/Fonts/STHeiti Medium.ttc", "STHeiti"),
    ]
    for path, name in candidates:
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            return name
        except Exception:
            continue
    # ????? reportlab ??? CID ????????????
    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        return 'STSong-Light'
    except Exception:
        return None


def plan_to_docx(plan: FullPlan) -> bytes:
    """将 FullPlan 转换为 Word 文档，返回 bytes。"""
    doc = Document()

    # 标题
    title = doc.add_heading(plan.input.course.name or "项目计划", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    p = doc.add_paragraph()
    p.add_run("项目要求：").bold = True
    p.add_run(plan.input.course.description or "无")
    p = doc.add_paragraph()
    p.add_run("截止日期：").bold = True
    p.add_run(str(plan.input.deadline))
    members_str = "、".join(
        f"{m.name}({m.daily_available_hours}h/天)" for m in plan.input.members
    )
    p = doc.add_paragraph()
    p.add_run("团队成员：").bold = True
    p.add_run(members_str or "无")

    doc.add_paragraph()  # 空行

    # 概述
    if plan.plan.summary:
        doc.add_heading("一、计划概述", level=1)
        _md_to_docx(doc, plan.plan.summary, base_heading_level=2)

    # 任务列表
    doc.add_heading("二、任务拆解", level=1)
    table = doc.add_table(rows=1, cols=6, style="Table Grid")
    hdr = table.rows[0].cells
    for i, text in enumerate(["ID", "任务", "工时", "依赖", "所需技能", "状态"]):
        hdr[i].text = text
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    status_map = {"pending": "待开始", "in_progress": "进行中",
                  "completed": "已完成", "blocked": "阻塞"}
    for t in plan.plan.tasks:
        row = table.add_row().cells
        row[0].text = t.id
        row[1].text = t.name
        row[2].text = f"{t.estimated_hours}h"
        row[3].text = ", ".join(t.dependencies) if t.dependencies else "-"
        row[4].text = ", ".join(t.required_skills) if t.required_skills else "-"
        row[5].text = status_map.get(t.status, t.status)

    # 时间线
    if plan.timeline.tasks:
        doc.add_heading("三、时间线安排", level=1)
        doc.add_paragraph(f"总工期 {plan.timeline.total_days} 天")
        tl_table = doc.add_table(rows=1, cols=5, style="Table Grid")
        hdr = tl_table.rows[0].cells
        for i, text in enumerate(["任务", "开始", "结束", "关键路径", "负责人"]):
            hdr[i].text = text
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for t in plan.timeline.tasks:
            row = tl_table.add_row().cells
            row[0].text = f"{t.task_id} {t.name}"
            row[1].text = str(t.start_date)
            row[2].text = str(t.end_date)
            row[3].text = "是" if t.is_critical else ""
            row[4].text = ", ".join(t.assigned_to) if t.assigned_to else "-"

    # QA 矩阵
    if plan.qa_matrix.assignments:
        doc.add_heading("四、责任分工", level=1)
        qa_table = doc.add_table(rows=1, cols=4, style="Table Grid")
        hdr = qa_table.rows[0].cells
        for i, text in enumerate(["任务", "负责人", "主要协助", "辅助协助"]):
            hdr[i].text = text
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for a in plan.qa_matrix.assignments:
            row = qa_table.add_row().cells
            row[0].text = a.task_name
            row[1].text = a.presenter
            row[2].text = a.qa_primary
            row[3].text = ", ".join(a.qa_support) if a.qa_support else "-"

    # 风险提示
    if plan.report.risk_note:
        doc.add_heading("五、风险提示", level=1)
        _md_to_docx(doc, plan.report.risk_note, base_heading_level=2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def plan_to_pdf(plan: FullPlan) -> bytes:
    """将 FullPlan 转换为 PDF 文档，返回 bytes。"""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    font_name = _register_cjk_font()
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
                                fontName=font_name or "Helvetica", fontSize=10,
                                leading=16)
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"],
                              fontName=font_name or "Helvetica", fontSize=16)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                              fontName=font_name or "Helvetica", fontSize=13)
    h3_style = ParagraphStyle("H3", parent=styles["Heading3"],
                              fontName=font_name or "Helvetica", fontSize=11)
    story = []

    story.append(Paragraph(saxutils.escape(plan.input.course.name) or "项目计划", h1_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>课程要求：</b>{saxutils.escape(plan.input.course.description or '无')}", body_style))
    story.append(Paragraph(f"<b>截止日期：</b>{plan.input.deadline}", body_style))
    members_str = "、".join(
        f"{m.name}({m.daily_available_hours}h/天)" for m in plan.input.members
    )
    story.append(Paragraph(f"<b>团队成员：</b>{saxutils.escape(members_str) or '无'}", body_style))
    story.append(Spacer(1, 16))

    if plan.plan.summary:
        story.append(Paragraph("一、计划概述", h2_style))
        story.extend(_md_to_pdf_story(plan.plan.summary, body_style, h2_style, h3_style, font_name))
        story.append(Spacer(1, 10))

    story.append(Paragraph("二、任务拆解", h2_style))
    status_map = {"pending": "待开始", "in_progress": "进行中",
                  "completed": "已完成", "blocked": "阻塞"}
    task_data = [["ID", "任务", "工时", "依赖", "状态"]]
    for t in plan.plan.tasks:
        task_data.append([
            t.id, t.name, f"{t.estimated_hours}h",
            ", ".join(t.dependencies) or "-", status_map.get(t.status, t.status),
        ])
    story.append(_build_table(task_data, font_name, body_style))
    story.append(Spacer(1, 10))

    if plan.timeline.tasks:
        story.append(Paragraph("三、时间线安排", h2_style))
        story.append(Paragraph(f"总工期 {plan.timeline.total_days} 天", body_style))
        tl_data = [["任务", "开始", "结束", "关键", "负责人"]]
        for t in plan.timeline.tasks:
            tl_data.append([
                f"{t.task_id} {t.name}", str(t.start_date), str(t.end_date),
                "是" if t.is_critical else "",
                ", ".join(t.assigned_to) if t.assigned_to else "-",
            ])
        story.append(_build_table(tl_data, font_name, body_style))
        story.append(Spacer(1, 10))

    if plan.qa_matrix.assignments:
        story.append(Paragraph("四、责任分工", h2_style))
        qa_data = [["任务", "负责人", "主要协助", "辅助协助"]]
        for a in plan.qa_matrix.assignments:
            qa_data.append([
                a.task_name, a.presenter, a.qa_primary,
                ", ".join(a.qa_support) if a.qa_support else "-",
            ])
        story.append(_build_table(qa_data, font_name, body_style))
        story.append(Spacer(1, 10))

    if plan.report.risk_note:
        story.append(Paragraph("五、风险提示", h2_style))
        risk_style = ParagraphStyle("Risk", parent=body_style, textColor=colors.red)
        story.extend(_md_to_pdf_story(plan.report.risk_note, risk_style, h2_style, h3_style, font_name))

    doc.build(story)
    return buf.getvalue()


def _build_table(data, font_name, style=None):
    """构造带样式的表格。"""
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph
    n_cols = len(data[0]) if data else 1
    # 按列数均分可用宽度（A4 可用约 17cm），避免溢出
    col_w = 17 * cm / n_cols
    # 用 Paragraph 包裹单元格，reportlab 自动换行
    wrapped = [[Paragraph(str(cell), style) if style else cell for cell in row] for row in data]
    t = Table(wrapped, repeatRows=1, colWidths=[col_w] * n_cols)
    style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4f46e5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
    ])
    if font_name:
        style.add("FONTNAME", (0, 0), (-1, -1), font_name)
    t.setStyle(style)
    return t
